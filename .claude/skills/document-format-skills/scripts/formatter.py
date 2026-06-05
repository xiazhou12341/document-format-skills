#!/usr/bin/env python3
"""
文档格式统一 v5
修复问题：
- 标题检测更全面
- 主送机关顶格
- 落款右对齐
- 清除斜体、下划线、颜色
- 特殊段落处理（附件、特此说明等）

公文标准：
- 页边距：上37mm，下35mm，左28mm，右26mm
- 主标题：居中，二号（22pt），方正小标宋简体
- 主送机关：顶格，三号仿宋
- 正文：三号仿宋GB2312，首行缩进2字符，行距28磅
- 一级标题："一、" 三号黑体，首行缩进2字符
- 二级标题："（一）" 三号楷体GB2312，首行缩进2字符
- 三级标题："1." 三号仿宋GB2312，首行缩进2字符
- 四级标题："（1）" 三号仿宋GB2312，首行缩进2字符
- 落款：右对齐，三号仿宋
- 附件：顶格，三号仿宋
"""

import sys
import re
import platform
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# macOS 字体映射：Windows 专属字体 → 方正 GBK 字体（用户已安装）
MAC_FONT_MAP = {
    '仿宋_GB2312': 'FZFangSong-Z02',
    '楷体_GB2312': 'FZKai-Z03',
    '方正小标宋简体': 'FZXiaoBiaoSong-B05',
    '黑体': 'FZHei-B01',
    '宋体': 'STSong',
    '仿宋': 'FZFangSong-Z02',
    '楷体': 'FZKai-Z03',
}

def resolve_font(font_name):
    """在 macOS 上将 Windows 字体映射为等价字体"""
    if platform.system() == 'Darwin':
        return MAC_FONT_MAP.get(font_name, font_name)
    return font_name

# 字号对照：二号=22pt，三号=16pt，小四=12pt
# 2字符缩进 = 2 × 16pt = 32pt（三号字）

# 自定义配置文件路径
def load_custom_preset():
    """加载自定义预设"""
    custom_config_file = Path(__file__).parent.parent / "custom_settings.json"
    if custom_config_file.exists():
        try:
            import json
            with open(custom_config_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            return None
    return None

PRESETS = {
    'official': {
        'name': '公文格式',
        'page': {'top': 3.7, 'bottom': 3.5, 'left': 2.8, 'right': 2.6},
        # 主标题：二号方正小标宋简体，加粗，居中
        'title': {
            'font_cn': '方正小标宋简体',
            'font_en': 'Times New Roman',
            'size': 22,  # 二号
            'bold': True,
            'align': 'center',
            'indent': 0,
        },
        # 主送机关：三号仿宋，顶格
        'recipient': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'left',
            'indent': 0,  # 顶格
        },
        # 一级标题：三号黑体，"一、"，首行缩进2字符
        'heading1': {
            'font_cn': '黑体',
            'font_en': 'Times New Roman',
            'size': 16,  # 三号
            'bold': False,
            'align': 'left',
            'indent': 32,  # 2字符缩进
        },
        # 二级标题：三号楷体GB2312，"（一）"，首行缩进2字符
        'heading2': {
            'font_cn': '楷体_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'left',
            'indent': 32,
        },
        # 三级标题：三号仿宋GB2312，"1."，首行缩进2字符
        'heading3': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'left',
            'indent': 32,
        },
        # 四级标题：三号仿宋GB2312，"（1）"，首行缩进2字符
        'heading4': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'left',
            'indent': 32,
        },
        # 正文：三号仿宋GB2312，首行缩进2字符（32pt），行距28磅
        'body': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'justify',
            'indent': 32,  # 2字符 = 2×16pt
            'line_spacing': 28,
        },
        # 落款单位：三号仿宋，右对齐
        'signature': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'right',
            'indent': 0,
        },
        # 落款日期：三号仿宋，右对齐
        'date': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'right',
            'indent': 0,
        },
        # 附件行：三号仿宋，格式同正文（首行缩进2字符）
        'attachment': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'justify',
            'indent': 32,  # 同正文，首行缩进2字符
        },
        # 结束语（特此说明/通知等）：三号仿宋，首行缩进
        'closing': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'left',
            'indent': 32,
        },
    },
    'academic': {
        'name': '学术论文格式',
        'page': {'top': 2.5, 'bottom': 2.5, 'left': 2.5, 'right': 2.5},
        'title': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 18, 'bold': True, 'align': 'center', 'indent': 0},
        'recipient': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'left', 'indent': 0},
        'heading1': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 15, 'bold': True, 'align': 'left', 'indent': 0},
        'heading2': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 14, 'bold': True, 'align': 'left', 'indent': 0},
        'heading3': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'left', 'indent': 0},
        'heading4': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'left', 'indent': 0},
        'body': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'justify', 'indent': 24, 'line_spacing': None},
        'signature': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'right', 'indent': 0},
        'date': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'right', 'indent': 0},
        'attachment': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'left', 'indent': 0},
        'closing': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'left', 'indent': 24},
    },
    'legal': {
        'name': '法律文书格式',
        'page': {'top': 3.0, 'bottom': 2.5, 'left': 3.0, 'right': 2.5},
        'title': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 22, 'bold': True, 'align': 'center', 'indent': 0},
        'recipient': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'heading1': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'heading2': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'heading3': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'heading4': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'body': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'justify', 'indent': 28, 'line_spacing': None},
        'signature': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'right', 'indent': 0},
        'date': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'right', 'indent': 0},
        'attachment': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'closing': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 28},
    },
}


def remove_background(doc):
    """移除页面背景颜色"""
    body = doc._body._body
    document = body.getparent()
    for elem in list(document):
        tag_name = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag_name == 'background':
            document.remove(elem)
    
    for para in doc.paragraphs:
        pPr = para._p.get_or_add_pPr()
        shd = pPr.find(qn('w:shd'))
        if shd is not None:
            pPr.remove(shd)
        for run in para.runs:
            run.font.highlight_color = None
            rPr = run._r.get_or_add_rPr()
            shd = rPr.find(qn('w:shd'))
            if shd is not None:
                rPr.remove(shd)


def _iter_block_items(doc):
    """Yield paragraphs and tables in document order."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith('}p'):
            yield Paragraph(child, doc)
        elif child.tag.endswith('}tbl'):
            yield Table(child, doc)


def _set_table_borders(table, size_pt=0.5, color="000000"):
    size = max(1, int(size_pt * 8))  # OOXML border size is in 1/8 pt
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    else:
        for child in list(borders):
            borders.remove(child)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), str(size))
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), color)
        borders.append(elem)


def _set_table_cell_margins(table, top_cm=0.0, bottom_cm=0.0, left_cm=0.05, right_cm=0.05):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    cell_mar = tbl_pr.find(qn('w:tblCellMar'))
    if cell_mar is None:
        cell_mar = OxmlElement('w:tblCellMar')
        tbl_pr.append(cell_mar)

    def _set_side(tag, cm_value):
        node = cell_mar.find(qn(f'w:{tag}'))
        if node is None:
            node = OxmlElement(f'w:{tag}')
            cell_mar.append(node)
        node.set(qn('w:type'), 'dxa')
        node.set(qn('w:w'), str(int(Cm(cm_value).twips)))

    _set_side('top', top_cm)
    _set_side('bottom', bottom_cm)
    _set_side('left', left_cm)
    _set_side('right', right_cm)


def _set_table_width_percent(table, percent=100):
    percent = max(1, min(100, int(percent)))
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:type'), 'pct')
    tbl_w.set(qn('w:w'), str(percent * 50))  # 50ths of a percent


def _set_table_indent(table, indent_twips=0):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:type'), 'dxa')
    tbl_ind.set(qn('w:w'), str(int(indent_twips)))


def _text_weight(text):
    weight = 0.0
    for ch in text:
        if ord(ch) < 128:
            weight += 0.5
        else:
            weight += 1.0
    return weight


def _normalize_pcts(weights, min_pct, max_pct):
    total = sum(weights) or 1.0
    pcts = [w / total * 100 for w in weights]

    # Clamp low
    for i, v in enumerate(pcts):
        if v < min_pct:
            pcts[i] = min_pct
    # Clamp high
    for i, v in enumerate(pcts):
        if v > max_pct:
            pcts[i] = max_pct

    # Renormalize to 100
    total = sum(pcts) or 1.0
    return [v / total * 100 for v in pcts]


def _set_table_col_widths_by_content(table, min_pct=8, max_pct=45):
    if not table.rows:
        return
    col_count = max(len(row.cells) for row in table.rows)
    if col_count == 0:
        return

    max_weights = [1.0] * col_count
    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            text = ''.join(p.text for p in cell.paragraphs).strip()
            if text:
                max_weights[c_idx] = max(max_weights[c_idx], _text_weight(text))

    pcts = _normalize_pcts(max_weights, min_pct, max_pct)

    # Set table grid + cell widths in pct
    tbl = table._tbl
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement('w:tblGrid')
        tbl.insert(0, tbl_grid)
    else:
        for child in list(tbl_grid):
            tbl_grid.remove(child)

    for pct in pcts:
        grid_col = OxmlElement('w:gridCol')
        grid_col.set(qn('w:w'), str(int(pct * 50)))  # pct in 1/50th %
        tbl_grid.append(grid_col)

    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            tc = cell._tc
            tc_pr = tc.tcPr
            if tc_pr is None:
                tc_pr = OxmlElement('w:tcPr')
                tc.insert(0, tc_pr)
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:type'), 'pct')
            tc_w.set(qn('w:w'), str(int(pcts[c_idx] * 50)))


def _insert_paragraph_after_table(table, text=""):
    p = OxmlElement("w:p")
    table._tbl.addnext(p)
    para = Paragraph(p, table._parent)
    if text:
        para.add_run(text)
    return para


def _insert_paragraph_before_table(table, text=""):
    p = OxmlElement("w:p")
    table._tbl.addprevious(p)
    para = Paragraph(p, table._parent)
    if text:
        para.add_run(text)
    return para


def _insert_paragraph_after_paragraph(paragraph, text=""):
    p = OxmlElement("w:p")
    paragraph._p.addnext(p)
    para = Paragraph(p, paragraph._parent)
    if text:
        para.add_run(text)
    return para


def _insert_paragraph_before_paragraph(paragraph, text=""):
    p = OxmlElement("w:p")
    paragraph._p.addprevious(p)
    para = Paragraph(p, paragraph._parent)
    if text:
        para.add_run(text)
    return para


def _is_numeric_text(text):
    text = text.replace(',', '').replace('％', '%').strip()
    if not text:
        return False
    import re
    return re.match(r'^[-+]?\\d+(?:\\.\\d+)?%?$', text) is not None


def _is_short_text(text, max_len=4):
    text = text.strip()
    return 0 < len(text) <= max_len


def _is_table_title(text):
    text = text.strip()
    if not text:
        return False
    if len(text) > 30:
        return False
    import re
    return re.match(r'^表\\s*(?:\\d+|[一二三四五六七八九十]+)(?:[\\-\\—\\._、]\\d+)?', text) is not None


def _is_table_unit(text):
    text = text.strip()
    if not text:
        return False
    if len(text) > 20:
        return False
    import re
    return re.match(r'^单位\\s*[:：]', text) is not None


def _set_cell_borders(cell, size_pt=0.5, color="000000"):
    size = max(1, int(size_pt * 8))
    tc = cell._tc
    tc_pr = tc.tcPr
    if tc_pr is None:
        tc_pr = OxmlElement('w:tcPr')
        tc.insert(0, tc_pr)

    borders = tc_pr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    else:
        for child in list(borders):
            borders.remove(child)

    for edge in ('top', 'left', 'bottom', 'right'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), str(size))
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), color)
        borders.append(elem)


def detect_para_type(text, index, total, alignment, all_texts):
    """
    检测段落类型
    返回: 'title', 'recipient', 'heading1', 'heading2', 'heading3', 'heading4', 
          'body', 'signature', 'date', 'attachment', 'closing'
    
    参数:
        text: 段落文本
        index: 段落索引
        total: 总段落数
        alignment: 原始对齐方式
        all_texts: 所有非空段落的文本列表，用于上下文判断
    """
    text = text.strip()
    if not text:
        return 'empty'
    
    # ===== 一级标题："一、" "二、" 等 =====
    if re.match(r'^[一二三四五六七八九十]+、', text):
        return 'heading1'
    
    # ===== 二级标题："（一）" "（二）" 等 =====
    if re.match(r'^（[一二三四五六七八九十]+）', text):
        return 'heading2'
    if re.match(r'^\([一二三四五六七八九十]+\)', text):
        return 'heading2'
    
    # ===== 三级标题："1." "2." 等 =====
    if re.match(r'^\d+\.\s*\S', text) and len(text) < 60:
        return 'heading3'
    
    # ===== 四级标题："（1）" "（2）" 等 =====
    if re.match(r'^（\d+）', text) and len(text) < 60:
        return 'heading4'
    if re.match(r'^\(\d+\)', text) and len(text) < 60:
        return 'heading4'
    
    # ===== 主送机关：XXX： 或 XXX: =====
    # 通常在文档开头几段，以冒号结尾，且较短
    if re.match(r'^[\u4e00-\u9fff]+[：:]$', text) and len(text) < 20:
        return 'recipient'
    
    # ===== 附件行 =====
    if re.match(r'^附件[：:]\s*', text):
        return 'attachment'
    if re.match(r'^附件\d*[：:．.\s]', text):
        return 'attachment'
    if re.match(r'^附件$', text):
        return 'attachment'
    
    # ===== 结束语 =====
    closing_patterns = [
        r'^特此(说明|通知|报告|函复|函告|批复|公告|通报)。?$',
        r'^此致$',
        r'^敬礼[！!]?$',
        r'^以上(报告|意见|方案).{0,10}$',
        r'^妥否.{0,10}$',
        r'^请.{0,15}(批示|审批|审议|指示|核准)。?$',
    ]
    for pattern in closing_patterns:
        if re.match(pattern, text):
            return 'closing'
    
    # ===== 落款日期 =====
    # 支持多种日期格式
    date_patterns = [
        r'^\d{4}年\d{1,2}月\d{1,2}日$',
        r'^\d{4}\.\d{1,2}\.\d{1,2}$',
        r'^\d{4}/\d{1,2}/\d{1,2}$',
        r'^\d{4}-\d{1,2}-\d{1,2}$',
        r'^二[○〇零oO0][一二三四五六七八九零〇○oO0]{2}年.{1,3}月.{1,3}日$',
    ]
    for pattern in date_patterns:
        if re.match(pattern, text):
            return 'date'
    
    # ===== 落款单位 =====
    # 判断逻辑：在文档后部，短文本，且下一段是日期或者是文档末尾附近
    if index >= total - 10 and len(text) < 30:
        # 检查是否像单位名称
        if re.search(r'(公司|局|委|部|厅|院|所|中心|办公室|集团|银行|学校|大学|医院)$', text):
            return 'signature'
        # 或者检查下文是否有日期
        remaining_texts = all_texts[all_texts.index(text)+1:] if text in all_texts else []
        for next_text in remaining_texts[:3]:
            for pattern in date_patterns:
                if re.match(pattern, next_text.strip()):
                    return 'signature'
    
    # ===== 主标题 =====
    # 判断条件：在前5段，且满足以下条件之一
    if index < 5:
        # 1. 明确的公文标题模式
        title_patterns = [
            r'^关于.+的(通知|报告|请示|函|意见|决定|公告|通报|批复|说明|方案|总结|汇报|复函|答复|建议)$',
            r'^.{2,30}(通知|报告|请示|函|意见|决定|公告|通报|批复|工作方案|工作总结|实施方案|管理办法|暂行规定)$',
        ]
        for pattern in title_patterns:
            if re.match(pattern, text):
                return 'title'
        
        # 2. 较长的标题（20-80字符），不以标点结尾
        if 15 < len(text) < 80 and not re.search(r'[。！？，、；：]$', text):
            # 排除以序号开头的
            if not re.match(r'^[一二三四五六七八九十\d（(]', text):
                return 'title'
        
        # 3. 居中的短文本（原本就是居中的）
        if alignment == WD_ALIGN_PARAGRAPH.CENTER and len(text) < 60:
            return 'title'
    
    # ===== 其他都是正文 =====
    return 'body'


def _split_heading_by_punct(paragraph):
    """Split heading like '（三）xxx：正文' or '（三）xxx\\n正文' into heading paragraph + body paragraph."""
    text = paragraph.text.strip()
    if not text:
        return False

    # Heading prefix patterns
    if not (
        re.match(r'^[一二三四五六七八九十]+、', text) or
        re.match(r'^（[一二三四五六七八九十]+）', text) or
        re.match(r'^\([一二三四五六七八九十]+\)', text) or
        re.match(r'^\d+\.\s*\S', text) or
        re.match(r'^（\d+）', text) or
        re.match(r'^\(\d+\)', text)
    ):
        return False

    # Find split positions:
    # 1. Soft-return (line break) is the primary separator
    # 2. Heading-ending punctuation (：, :, 。) are secondary separators
    punct_positions = []

    # Line breaks (soft return) separate heading from body text
    if '\n' in text:
        punct_positions.append(text.find('\n'))

    # Only use ：, :, 。 as heading-ending punctuation
    for ch in ('：', ':', '。'):
        pos = text.find(ch)
        if pos != -1:
            punct_positions.append(pos)

    if not punct_positions:
        return False
    split_idx = min(punct_positions)

    # For line break, don't include the \\n in the heading text
    if text[split_idx] == '\n':
        head = text[:split_idx].strip()
    else:
        head = text[:split_idx + 1].strip()
    tail = text[split_idx + 1:].strip()
    if not tail:
        return False

    paragraph.text = head
    new_para = _insert_paragraph_after_paragraph(paragraph, text=tail)
    return new_para is not None


def set_font(run, font_cn, font_en, size, bold=False):
    """
    设置字体，同时清除原有格式（斜体、下划线、颜色）
    macOS 上自动将 Windows 字体映射为等价字体
    """
    # macOS 字体兼容
    font_cn = resolve_font(font_cn)
    font_en = resolve_font(font_en)

    # 基本字体设置
    run.font.name = font_en
    run.font.size = Pt(size)
    run.font.bold = bold
    
    # 清除斜体
    run.font.italic = False
    
    # 清除下划线
    run.font.underline = False
    
    # 清除颜色（设置为黑色）
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # 清除删除线
    run.font.strike = False
    run.font.double_strike = False
    
    # 清除上下标
    run.font.subscript = False
    run.font.superscript = False
    
    # 设置中文字体
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)
    rFonts.set(qn('w:cs'), font_en)


def format_paragraph(para, fmt, para_type, line_spacing_pt=28, first_line_bold=False):
    """格式化段落
    
    fmt 支持的字段:
        font_cn, font_en, size, bold, align, indent,
        line_spacing  - 行距(磅), None表示使用1.5倍行距
        space_before  - 段前间距(磅), 默认0
        space_after   - 段后间距(磅), 默认0
    """
    pf = para.paragraph_format
    
    # 对齐方式
    align_map = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    pf.alignment = align_map.get(fmt.get('align', 'justify'), WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    # 段落左缩进清零（重要：确保"文本之前缩进"为0）
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    
    # 首行缩进
    indent = fmt.get('indent', 0)
    if indent > 0:
        pf.first_line_indent = Pt(indent)
    else:
        pf.first_line_indent = Pt(0)
    
    # 行距：优先读取当前元素自身的 line_spacing，否则用全局默认值
    ls = fmt.get('line_spacing', line_spacing_pt)
    if ls:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(ls)
    else:
        pf.line_spacing = 1.5
    
    # 段前段后（支持自定义，默认0）
    pf.space_before = Pt(fmt.get('space_before', 0))
    pf.space_after = Pt(fmt.get('space_after', 0))
    
    # 字体 - 支持首句加粗
    if first_line_bold and para_type == 'body':
        # 首句以中文句号“。”作为结束
        full_text = para.text
        first_sentence_end = full_text.find('。')
        if first_sentence_end != -1:
            split_idx = first_sentence_end + 1
            first_part = full_text[:split_idx]
            rest_part = full_text[split_idx:]
            
            # 重新构建 runs，确保只加粗首句
            for run in list(para.runs):
                para._p.remove(run._r)
            
            run1 = para.add_run(first_part)
            set_font(run1, fmt['font_cn'], fmt['font_en'], fmt['size'], bold=True)
            
            if rest_part:
                run2 = para.add_run(rest_part)
                set_font(run2, fmt['font_cn'], fmt['font_en'], fmt['size'], fmt.get('bold', False))
        else:
            # 没找到中文句号，正常处理
            for run in para.runs:
                set_font(run, fmt['font_cn'], fmt['font_en'], fmt['size'], fmt.get('bold', False))
    else:
        # 正文里的“一是/二是...”加粗前缀
        if para_type == 'body':
            import re as _re
            m = _re.match(r'^([一二三四五六七八九十]{1,3}是)([：:、]?)', para.text)
            if m:
                lead = m.group(1) + (m.group(2) or '')
                rest = para.text[len(lead):]
                for run in list(para.runs):
                    para._p.remove(run._r)
                run1 = para.add_run(lead)
                set_font(run1, fmt['font_cn'], fmt['font_en'], fmt['size'], bold=True)
                if rest:
                    run2 = para.add_run(rest)
                    set_font(run2, fmt['font_cn'], fmt['font_en'], fmt['size'], fmt.get('bold', False))
                return

        # 正常处理
        for run in para.runs:
            set_font(run, fmt['font_cn'], fmt['font_en'], fmt['size'], fmt.get('bold', False))


def add_page_number(doc, font_name="宋体"):
    """添加页码（四号宋体，左右一字线，奇右偶左，距版心下边缘约7mm）"""
    # 启用奇偶页页眉页脚（文档级）
    try:
        doc.settings.odd_and_even_pages_header_footer = True
    except Exception:
        settings_el = doc.settings._element
        if settings_el.find(qn('w:evenAndOddHeaders')) is None:
            settings_el.append(OxmlElement('w:evenAndOddHeaders'))

    for section in doc.sections:

        section.odd_and_even_pages_header_footer = True
        section.footer_distance = Cm(0.7)

        odd_footer = section.footer
        even_footer = section.even_page_footer
        odd_footer.is_linked_to_previous = False
        even_footer.is_linked_to_previous = False

        for f in (odd_footer, even_footer):
            for para in f.paragraphs:
                para.clear()

        def _build_footer_line(footer, align, pad_fullwidth):
            if footer.paragraphs:
                para = footer.paragraphs[0]
            else:
                para = footer.add_paragraph()

            para.alignment = align

            # 前导全角空格（空一字）
            if pad_fullwidth:
                run0 = para.add_run("　")
                set_font(run0, font_name, font_name, 14, bold=False)

            # 左一字线（带空格）
            run1 = para.add_run("— ")
            set_font(run1, font_name, font_name, 14, bold=False)

            # 页码域
            run2 = para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run2._r.append(fldChar1)
            set_font(run2, font_name, font_name, 14, bold=False)

            run3 = para.add_run()
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            run3._r.append(instrText)
            set_font(run3, font_name, font_name, 14, bold=False)

            run4 = para.add_run()
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run4._r.append(fldChar2)
            set_font(run4, font_name, font_name, 14, bold=False)

            # 右一字线（带空格）
            run5 = para.add_run(" —")
            set_font(run5, font_name, font_name, 14, bold=False)

            # 末尾全角空格（空一字）
            if not pad_fullwidth:
                run6 = para.add_run("　")
                set_font(run6, font_name, font_name, 14, bold=False)

        # 奇数页居右空一字，偶数页居左空一字
        _build_footer_line(odd_footer, WD_ALIGN_PARAGRAPH.RIGHT, pad_fullwidth=True)
        _build_footer_line(even_footer, WD_ALIGN_PARAGRAPH.LEFT, pad_fullwidth=False)


def format_document(input_path, output_path, preset_name='official'):
    """格式化文档"""
    # 处理自定义预设
    if preset_name == 'custom':
        preset = load_custom_preset()
        if preset is None:
            print('Custom preset not found, using official preset')
            preset = PRESETS['official']
        else:
            print(f'Preset: {preset.get("name", "自定义格式")}')
    elif preset_name not in PRESETS:
        print(f'Unknown preset: {preset_name}')
        print(f'Available: {", ".join(PRESETS.keys())}')
        sys.exit(1)
    else:
        preset = PRESETS[preset_name]
        print(f'Preset: {preset["name"]}')
    
    print(f'Input: {input_path}')
    
    # 获取首句加粗选项
    first_line_bold = preset.get('first_line_bold', False)
    
    doc = Document(input_path)

    # 将“标题+标点+正文”拆分为标题段+正文段
    for para in list(doc.paragraphs):
        _split_heading_by_punct(para)

    total_paras = len(doc.paragraphs)
    
    # 收集所有非空段落文本，用于上下文判断
    all_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    # 1. 移除背景
    print('1. Removing background...')
    remove_background(doc)
    
    # 2. 设置页面边距
    print('2. Setting page margins...')
    page = preset['page']
    for section in doc.sections:
        section.top_margin = Cm(page['top'])
        section.bottom_margin = Cm(page['bottom'])
        section.left_margin = Cm(page['left'])
        section.right_margin = Cm(page['right'])
    
    # 3. 格式化段落
    print('3. Formatting paragraphs...')
    stats = {
        'title': 0, 'recipient': 0, 'heading1': 0, 'heading2': 0, 
        'heading3': 0, 'heading4': 0, 'body': 0, 'signature': 0, 
        'date': 0, 'attachment': 0, 'closing': 0
    }
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        para_type = detect_para_type(
            text, i, total_paras, 
            para.paragraph_format.alignment,
            all_texts
        )
        
        # 选择对应的格式
        fmt_key = para_type if para_type in preset else 'body'
        fmt = preset.get(fmt_key, preset['body'])
        
        format_paragraph(para, fmt, para_type, first_line_bold=first_line_bold)
        stats[para_type] = stats.get(para_type, 0) + 1
        
        # 打印处理信息
        preview = text[:35] + '...' if len(text) > 35 else text
        print(f'   [{para_type:10}] {preview}')
    
    # 4. 处理表格
    print('4. Formatting tables...')
    body_fmt = preset.get('body', {})
    # 表格配置：优先使用 preset 中的 table 节点，否则用 body 格式
    table_fmt = preset.get('table', {})
    table_defaults = {
        'optimize': True,
        'border_size_pt': 0.5,
        'width_percent': 100,
        'auto_col_width': True,
        'col_min_pct': 8,
        'col_max_pct': 45,
        'row_height_cm': 0.7,
        'cell_margin_top_cm': 0.0,
        'cell_margin_bottom_cm': 0.0,
        'cell_margin_left_cm': 0.05,
        'cell_margin_right_cm': 0.05,
        'paragraph_single': True,
        'after_table_blank_line': True,
        'title_align': 'center',
        'unit_align': 'right',
        'unit_space_before_lines': 0.5,
        'short_text_len': 4,
    }
    table_cfg = {**table_defaults, **table_fmt}

    tbl_font_cn = table_fmt.get('font_cn', body_fmt.get('font_cn', '仿宋_GB2312'))
    tbl_font_en = table_fmt.get('font_en', body_fmt.get('font_en', 'Times New Roman'))
    tbl_size = table_fmt.get('size', body_fmt.get('size', 16))
    tbl_bold = table_fmt.get('bold', False)
    tbl_line_spacing = table_fmt.get('line_spacing', body_fmt.get('line_spacing', 28))
    tbl_header_bold = table_fmt.get('header_bold', False)
    tbl_first_line_indent = table_fmt.get('first_line_indent', 0)

    blocks = list(_iter_block_items(doc))
    for idx, block in enumerate(blocks):
        if not isinstance(block, Table):
            continue

        table = block
        if table_cfg.get('optimize', True):
            table.autofit = not table_cfg.get('auto_col_width', True)
            _set_table_width_percent(table, table_cfg.get('width_percent', 100))
            _set_table_indent(table, 0)
            _set_table_borders(table, size_pt=table_cfg.get('border_size_pt', 0.5))
            _set_table_cell_margins(
                table,
                top_cm=table_cfg.get('cell_margin_top_cm', 0.0),
                bottom_cm=table_cfg.get('cell_margin_bottom_cm', 0.0),
                left_cm=table_cfg.get('cell_margin_left_cm', 0.05),
                right_cm=table_cfg.get('cell_margin_right_cm', 0.05),
            )
            if table_cfg.get('auto_col_width', True):
                _set_table_col_widths_by_content(
                    table,
                    min_pct=table_cfg.get('col_min_pct', 8),
                    max_pct=table_cfg.get('col_max_pct', 45),
                )

        # 表格前空一行（如果已有空行则不重复）
        prev_block = blocks[idx - 1] if idx - 1 >= 0 else None
        prev_para_is_title = isinstance(prev_block, Paragraph) and _is_table_title(prev_block.text)
        prev_para_is_unit = isinstance(prev_block, Paragraph) and _is_table_unit(prev_block.text)

        if isinstance(prev_block, Paragraph):
            if prev_block.text.strip():
                if prev_para_is_title or prev_para_is_unit:
                    _insert_paragraph_before_paragraph(prev_block, text="")
                else:
                    _insert_paragraph_before_table(table, text="")
        elif isinstance(prev_block, Table):
            _insert_paragraph_after_table(prev_block, text="")
        else:
            if idx == 0:
                _insert_paragraph_before_table(table, text="")

        # 标题段落（表格前一段）
        if prev_para_is_title:
            if table_cfg.get('title_align', 'center') == 'center':
                prev_block.alignment = WD_ALIGN_PARAGRAPH.CENTER
            prev_block.paragraph_format.space_before = Pt(0)
            prev_block.paragraph_format.space_after = Pt(0)
            prev_block.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 单位段落（表格后一段）
        next_block = blocks[idx + 1] if idx + 1 < len(blocks) else None
        unit_para = None
        if isinstance(next_block, Paragraph) and _is_table_unit(next_block.text):
            unit_para = next_block
            if table_cfg.get('unit_align', 'right') == 'right':
                unit_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            unit_space_lines = table_cfg.get('unit_space_before_lines', 0.5)
            unit_para.paragraph_format.space_before = Pt(tbl_size * unit_space_lines)
            unit_para.paragraph_format.space_after = Pt(0)
            unit_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 表格内内容
        serial_col_idx = None
        if table.rows:
            header_cells = table.rows[0].cells
            for c_idx, cell in enumerate(header_cells):
                head_text = ''.join(p.text for p in cell.paragraphs).strip()
                if '序号' in head_text or head_text == '序':
                    serial_col_idx = c_idx
                    break

        for row_idx, row in enumerate(table.rows):
            # 行高
            if table_cfg.get('row_height_cm'):
                row.height = Cm(table_cfg.get('row_height_cm'))
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

            for col_idx, cell in enumerate(row.cells):
                if table_cfg.get('optimize', True):
                    _set_cell_borders(cell, size_pt=table_cfg.get('border_size_pt', 0.5))

                cell_text = ''.join(p.text for p in cell.paragraphs).strip()
                for para in cell.paragraphs:
                    # 字体设置
                    if para.text.strip():
                        is_header = (row_idx == 0 and tbl_header_bold)
                        for run in para.runs:
                            set_font(run, tbl_font_cn, tbl_font_en, tbl_size, bold=(tbl_bold or is_header))

                    # 段落格式
                    para.paragraph_format.first_line_indent = Pt(tbl_first_line_indent)
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    if table_cfg.get('paragraph_single', True):
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    else:
                        if tbl_line_spacing:
                            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                            para.paragraph_format.line_spacing = Pt(tbl_line_spacing)
                        else:
                            para.paragraph_format.line_spacing = 1.5

                    # 对齐策略
                    if row_idx == 0:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif '合计' in cell_text or '总计' in cell_text:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif serial_col_idx is not None and col_idx == serial_col_idx:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif _is_numeric_text(cell_text):
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif _is_short_text(cell_text, table_cfg.get('short_text_len', 4)):
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 表格后空一行（若已有空行则不重复）
        if table_cfg.get('after_table_blank_line', True):
            next_block = blocks[idx + 1] if idx + 1 < len(blocks) else None
            if unit_para is not None:
                # 单位行后再空一行
                after_unit = blocks[idx + 2] if idx + 2 < len(blocks) else None
                if not (isinstance(after_unit, Paragraph) and not after_unit.text.strip()):
                    _insert_paragraph_after_paragraph(unit_para, text="")
            else:
                if not (isinstance(next_block, Paragraph) and not next_block.text.strip()):
                    _insert_paragraph_after_table(table, text="")
    
    # 5. 添加页码
    if preset.get('page_number', True):
        print('5. Adding page numbers...')
        add_page_number(doc, font_name=preset.get('page_number_font', '宋体'))
    else:
        print('5. Skipping page numbers...')
    
    # 保存
    doc.save(output_path)
    
    print()
    print('=' * 50)
    print('Statistics:')
    for k, v in stats.items():
        if v > 0:
            print(f'  {k}: {v}')
    print(f'Output: {output_path}')


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: python formatter.py input.docx [output.docx] [--preset official|academic|legal]')
        print('  If output.docx is omitted, auto-generates as "{input}-已格式化.docx"')
        sys.exit(1)

    input_file = sys.argv[1]
    # Auto-generate output filename with "-已格式化" suffix
    if len(sys.argv) >= 3 and not sys.argv[2].startswith('--'):
        output_file = sys.argv[2]
    else:
        # Generate output name: input name + "-已格式化"
        p = Path(input_file)
        output_file = str(p.with_stem(p.stem + '-已格式化'))
    
    preset = 'official'
    if '--preset' in sys.argv:
        idx = sys.argv.index('--preset')
        if idx + 1 < len(sys.argv):
            preset = sys.argv[idx + 1]
    
    format_document(input_file, output_file, preset)
